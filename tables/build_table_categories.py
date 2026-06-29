import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Table2_categories_subdiagnoses.docx")
FONT, SIZE = "Arial", 12

HEADER = ["DSM category", "Diagnoses (n)", "Constituent diagnoses"]
ROWS = [
    ("Anxiety", "Generalized anxiety disorder; separation anxiety disorder; social anxiety disorder; "
     "panic disorder; agoraphobia; specific phobia"),
    ("Eating", "Anorexia nervosa; bulimia nervosa; binge-eating disorder"),
    ("Psychosis", "Schizophrenia; schizoaffective disorder; schizophreniform disorder"),
    ("Tic", "Tourette's disorder; persistent tic disorder; provisional tic disorder"),
    ("Depression", "Major depressive disorder; persistent depressive disorder"),
    ("Bipolar", "Bipolar I disorder; bipolar II disorder"),
    ("ADHD", "Attention-deficit/hyperactivity disorder"),
    ("ODD", "Oppositional defiant disorder"),
    ("Conduct", "Conduct disorder"),
    ("DMDD", "Disruptive mood dysregulation disorder"),
    ("OCD", "Obsessive-compulsive disorder"),
    ("PTSD", "Post-traumatic stress disorder"),
    ("Autism", "Autism spectrum disorder"),
]
NOTE = ("Category caseness tables aggregate the constituent diagnoses listed here; the sub-disorder "
        "caseness tables report them separately, so construct-membership choices (e.g., whether specific "
        "phobia is counted within anxiety) can be made directly. Other specified and unspecified "
        "(subthreshold) variants of each diagnosis are retained in the resolved layer and excluded from "
        "caseness by default. Sleep, suicidality, and homicidality are recorded in the resolved layer but "
        "are not summarized as caseness.")


def style_run(run, *, italic=False, bold=False):
    run.font.name = FONT
    run.font.size = Pt(SIZE)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.italic = italic
    run.bold = bold


def para(doc, parts, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    for text, kw in parts:
        style_run(p.add_run(text), **kw)
    return p


def set_border(cell, edge, sz=8):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders"); tcPr.append(borders)
    el = borders.find(qn(f"w:{edge}"))
    if el is None:
        el = OxmlElement(f"w:{edge}"); borders.append(el)
    el.set(qn("w:val"), "single"); el.set(qn("w:sz"), str(sz)); el.set(qn("w:color"), "000000")


def fill_cell(cell, text, *, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    style_run(p.add_run(text))


doc = Document()
sec = doc.sections[0]
sec.left_margin = sec.right_margin = Inches(1.0)
doc.styles["Normal"].font.name = FONT
doc.styles["Normal"].font.size = Pt(SIZE)

para(doc, [("Table 2", dict(bold=True))])
para(doc, [("DSM categories and their constituent KSADS-COMP diagnoses", dict(italic=True))], space_after=6)

tbl = doc.add_table(rows=len(ROWS) + 1, cols=3)
tbl.autofit = False
widths = [Inches(1.4), Inches(1.0), Inches(4.1)]
for row in tbl.rows:
    for j, w in enumerate(widths):
        row.cells[j].width = w
for j, h in enumerate(HEADER):
    fill_cell(tbl.rows[0].cells[j], h, center=(j != 0))
for i, (cat, dis) in enumerate(ROWS, start=1):
    n = dis.count(";") + 1
    fill_cell(tbl.rows[i].cells[0], cat)
    fill_cell(tbl.rows[i].cells[1], str(n), center=True)
    fill_cell(tbl.rows[i].cells[2], dis)
for c in tbl.rows[0].cells:
    set_border(c, "top"); set_border(c, "bottom")
for c in tbl.rows[-1].cells:
    set_border(c, "bottom")

para(doc, [("Note. ", dict(italic=True)), (NOTE, dict())])
doc.save(OUT)
print("wrote", OUT)