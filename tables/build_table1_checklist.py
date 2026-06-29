import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Table1_reporting_checklist.docx")
FONT, SIZE = "Arial", 12

HEADER = ["Analytic decision", "What to report", "Recommended default",
          "Include in sensitivity analysis?"]
ROWS = [
    ("Diagnostic status", "Current vs. ever-met (present, past, remission)",
     "Match to the research question", "Yes"),
    ("Informant", "Parent, youth, either, or both", "Either", "Yes"),
    ("Threshold", "Full criteria vs. including subthreshold", "Full criteria", "Optional"),
    ("Construct membership", "e.g., specific phobia in or out of anxiety",
     "State explicitly", "Yes, for anxiety"),
    ("Lifetime reconstruction", "Whether ever-met is computed within a wave or unioned across waves",
     "Union across waves for lifetime status", "Yes, for longitudinal analyses"),
    ("Administrative codes", "How 555, 888, and missing values were handled",
     "Not-administered kept as missing; never recoded to 0", "Fixed (correctness, not preference)"),
    ("Instrument version", "Which KSADS-COMP versions are included",
     "Report version; flag the 1.0-to-2.0 transition", "If the analysis spans ses-03A"),
    ("Administration schedule", "Waves and modules used",
     "Follow the administration calendar", "Not applicable"),
]
NOTE = ("The four upper rows are the analytic choices that vary across the specification "
        "curve in Figure 1; investigators should state each and report whether substantive "
        "results hold across reasonable alternatives. The recommended default is this "
        "resource's configuration, not an external standard. Administrative-code handling is "
        "fixed because it has a correct answer rather than a preference: a not-administered "
        "cell is not a negative. Within-wave episode fields do not carry prior diagnoses "
        "forward, so lifetime status should be reconstructed by unioning episodes across waves "
        "(Figure 6).")


def style_run(run, *, italic=False, bold=False):
    run.font.name = FONT
    run.font.size = Pt(SIZE)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.italic = italic
    run.bold = bold


def para(doc, parts, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    for text, kw in parts:
        style_run(p.add_run(text), **kw)
    return p


def set_border(cell, edge, sz=8):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    el = borders.find(qn(f"w:{edge}"))
    if el is None:
        el = OxmlElement(f"w:{edge}")
        borders.append(el)
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:color"), "000000")


def fill_cell(cell, text, *, bold=False, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    style_run(p.add_run(text), bold=bold)


doc = Document()
sec = doc.sections[0]
sec.left_margin = sec.right_margin = Inches(1.0)
for st in ("Normal",):
    doc.styles[st].font.name = FONT
    doc.styles[st].font.size = Pt(SIZE)

para(doc, [("Table 1", dict(bold=True))], space_after=0)
para(doc, [("Reporting checklist for KSADS-COMP diagnostic definitions", dict(italic=True))],
     space_after=6)

n = len(ROWS) + 1
tbl = doc.add_table(rows=n, cols=4)
tbl.autofit = False
widths = [Inches(1.5), Inches(2.2), Inches(2.0), Inches(1.6)]
for j, w in enumerate(widths):
    for row in tbl.rows:
        row.cells[j].width = w

for j, h in enumerate(HEADER):
    fill_cell(tbl.rows[0].cells[j], h, bold=False, center=True)
for i, r in enumerate(ROWS, start=1):
    for j, val in enumerate(r):
        fill_cell(tbl.rows[i].cells[j], val, center=False)

for c in tbl.rows[0].cells:
    set_border(c, "top")
    set_border(c, "bottom")
for c in tbl.rows[-1].cells:
    set_border(c, "bottom")

para(doc, [("Note. ", dict(italic=True)), (NOTE, dict())], space_after=0)
doc.save(OUT)
print("wrote", OUT)